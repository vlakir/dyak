"""Тесты контекста, имени файла, автопочинки пробелов и рендера docx (T006)."""

from __future__ import annotations

import logging
from pathlib import Path

import pytest
from docx import Document

from dyak.domain import Person
from dyak.errors import TemplateError
from dyak.render.context import build_context
from dyak.render.engine import (
    default_filename_template,
    fix_jinja_spaces,
    render_document,
    render_filename,
    reset_tag_warnings,
)
from dyak.render.filters import EMPTY_MARKER


def _person() -> Person:
    return Person(
        cells={
            'Фамилия': 'Иванов',
            'Имя': 'Пётр',
            'Отчество': 'Семёнович',
            'Должность': 'директор',
            'Дата_начала': '01.07.2026',
        },
    )


def test_build_context_is_flat_by_headers() -> None:
    ctx = build_context(_person())
    assert ctx['Фамилия'] == 'Иванов'
    assert ctx['Должность'] == 'директор'
    assert ctx['Дата_начала'] == '01.07.2026'


def test_build_context_splits_fullname_into_parts() -> None:
    person = Person(cells={'ФИО': 'Иванов Пётр Семёнович', 'Должность': 'директор'})
    ctx = build_context(person, fullname_source='ФИО')
    assert ctx['ФИО'] == 'Иванов Пётр Семёнович'  # целое остаётся доступным
    assert ctx['Фамилия'] == 'Иванов'
    assert ctx['Имя'] == 'Пётр'
    assert ctx['Отчество'] == 'Семёнович'


def test_build_context_fullname_two_words_has_no_patronymic() -> None:
    person = Person(cells={'ФИО': 'Петрова Анна'})
    ctx = build_context(person, fullname_source='ФИО')
    assert ctx['Фамилия'] == 'Петрова'
    assert ctx['Имя'] == 'Анна'
    assert ctx['Отчество'] == ''


def test_render_filename_substitutes_by_header() -> None:
    ctx = build_context(_person())
    assert render_filename('Приказ_{{ Фамилия }}.docx', ctx) == 'Приказ_Иванов.docx'


def test_render_filename_fixes_forgotten_space(
    caplog: pytest.LogCaptureFixture,
) -> None:
    ctx = build_context(_person())
    with caplog.at_level(logging.WARNING):
        name = render_filename('{{ Дата начала }}.docx', ctx)
    assert name == '01.07.2026.docx'
    assert any('Дата начала' in r.message for r in caplog.records)


def test_fix_jinja_spaces_leaves_filter_expr_untouched() -> None:
    # Тег с фильтром (форма T002) не трогается — пробелы там значимы.
    assert fix_jinja_spaces('{{ ФИО | рд }}') == '{{ ФИО | рд }}'
    # Управляющий блок тоже не трогается.
    assert fix_jinja_spaces('{% for x in y %}') == '{% for x in y %}'


def test_fix_jinja_spaces_normalizes_special_chars() -> None:
    # T026 issue 1: спецсимволы в «голом» теге → `_` (как в ключе контекста).
    assert fix_jinja_spaces('{{ л/н }}') == '{{ л_н }}'
    assert fix_jinja_spaces('{{ № п/п }}') == '{{ п_п }}'
    # Тег с атрибутом/вызовом не трогаем (символы значимы).
    assert fix_jinja_spaces('{{ ФИО.инициалы }}') == '{{ ФИО.инициалы }}'
    assert fix_jinja_spaces("{{ [a, b] | join('_') }}") == "{{ [a, b] | join('_') }}"


def test_fix_jinja_spaces_warns_once_per_tag(caplog: pytest.LogCaptureFixture) -> None:
    # T026: предупреждение авто-фикса — раз на уникальный тег за прогон (шаблон
    # патчится построчно, иначе на N строк было бы N повторов).
    reset_tag_warnings()
    with caplog.at_level(logging.WARNING):
        for _ in range(3):
            fix_jinja_spaces('{{ л/н }}')
    assert sum('л/н' in r.message for r in caplog.records) == 1
    reset_tag_warnings()  # новый прогон — предупреждаем снова
    with caplog.at_level(logging.WARNING):
        fix_jinja_spaces('{{ л/н }}')
    second_run = sum('л/н' in r.message for r in caplog.records)
    assert second_run == 2  # первый прогон (1) + второй после reset (1)


def test_default_filename_template_from_roles() -> None:
    roles = {'surname': 'Фамилия', 'name': 'Имя', 'patronymic': 'Отчество'}
    tpl = default_filename_template(roles)
    assert tpl is not None
    assert render_filename(tpl, build_context(_person())) == (
        'Иванов_Пётр_Семёнович.docx'
    )


def test_default_filename_drops_empty_parts() -> None:
    roles = {'surname': 'Фамилия', 'name': 'Имя', 'patronymic': 'Отчество'}
    tpl = default_filename_template(roles)
    assert tpl is not None
    person = Person(cells={'Фамилия': 'Петрова', 'Имя': 'Анна', 'Отчество': ''})
    # Пустое отчество не оставляет хвостового подчёркивания.
    assert render_filename(tpl, build_context(person)) == 'Петрова_Анна.docx'


def test_default_filename_template_none_without_roles() -> None:
    assert default_filename_template({}) is None


def test_render_document_substitutes_into_docx(tmp_path: Path) -> None:
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('Назначить {{ Фамилия }} {{ Имя }}')
    doc.add_paragraph('на {{ Должность }} с {{ Дата_начала }}.')
    doc.save(template)

    out = tmp_path / 'out.docx'
    render_document(template, build_context(_person()), out)

    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'Назначить Иванов Пётр' in text
    assert 'на директор с 01.07.2026.' in text


def test_render_document_collapses_double_space_from_empty_field(
    tmp_path: Path,
) -> None:
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('Назначить {{ Фамилия }} {{ Имя }} {{ Отчество }} (готово)')
    doc.save(template)

    # Отчество пустое — между «Анна» и «(готово)» не должно остаться двойного
    # пробела.
    person = Person(cells={'Фамилия': 'Петрова', 'Имя': 'Анна', 'Отчество': ''})
    out = tmp_path / 'out.docx'
    render_document(template, build_context(person, fullname_source=None), out)

    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'Назначить Петрова Анна (готово)' in text
    assert '  ' not in text


def test_render_document_collapses_spaces_in_table_cells(tmp_path: Path) -> None:
    template = tmp_path / 'tpl.docx'
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run('{{ Имя }} {{ Отчество }} конец')
    doc.save(template)

    person = Person(cells={'Имя': 'Анна', 'Отчество': ''})
    out = tmp_path / 'out.docx'
    render_document(template, build_context(person), out)

    cell_text = Document(out).tables[0].cell(0, 0).text
    assert cell_text == 'Анна конец'


def test_render_document_fixes_run_split_tag_with_space(
    tmp_path: Path,
    caplog: pytest.LogCaptureFixture,
) -> None:
    # Плейсхолдер с пробелом, разрезанный Word по run'ам, должен и склеиться,
    # и починиться (пробел → подчёркивание) + warning.
    template = tmp_path / 'tpl.docx'
    doc = Document()
    para = doc.add_paragraph()
    para.add_run('{{ Дата ')
    para.add_run('начала }}')
    doc.save(template)

    out = tmp_path / 'out.docx'
    with caplog.at_level(logging.WARNING):
        render_document(template, build_context(_person()), out)
    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert '01.07.2026' in text
    assert any('Дата начала' in r.message for r in caplog.records)


def test_render_document_raises_on_undefined_variable(tmp_path: Path) -> None:
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('{{ НетТакого }}')
    doc.save(template)
    out = tmp_path / 'out.docx'
    with pytest.raises(TemplateError, match='НетТакого'):
        render_document(template, {'Фамилия': 'Иванов'}, out)
    assert not out.exists()


def test_render_filename_raises_on_undefined_variable() -> None:
    with pytest.raises(TemplateError, match='НетТакого'):
        render_filename('{{ НетТакого }}.docx', {'Фамилия': 'Иванов'})


def test_render_document_wraps_syntax_error_friendly(tmp_path: Path) -> None:
    # T026 issue 3: грубая ошибка разметки → доменная TemplateError (понятный
    # текст в лог), а не сырой jinja-TemplateSyntaxError (трейсбек во всплывашке).
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('{{ Фамилия | }}')  # пустой фильтр — синтаксическая ошибка
    doc.save(template)
    with pytest.raises(TemplateError, match='разметке шаблона'):
        render_document(template, {'Фамилия': 'Иванов'}, tmp_path / 'out.docx')


def test_render_filename_wraps_syntax_error_friendly() -> None:
    with pytest.raises(TemplateError, match='разметке шаблона'):
        render_filename('{{ Фамилия | }}.docx', {'Фамилия': 'Иванов'})


# --- T016 фаза C: empty-aware чистка пробелов и висячей пунктуации -------------


def _render_body(tmp_path: Path, body: str, cells: dict[str, str]) -> str:
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph(body)
    doc.save(template)
    out = tmp_path / 'out.docx'
    render_document(template, build_context(Person(cells=cells)), out)
    return Document(out).paragraphs[0].text


def test_empty_value_drops_dangling_terminator_keeps_label(tmp_path: Path) -> None:
    # «Звание: ▮.» → «Звание:» (висячая точка убрана, лейбл остаётся).
    assert _render_body(tmp_path, 'Звание: {{ Звание }}.', {'Звание': ''}) == 'Звание:'


def test_empty_value_keeps_separator_comma(tmp_path: Path) -> None:
    # Запятая-разделитель сохраняется, пробел перед ней убирается.
    text = _render_body(tmp_path, 'Оклад {{ Премия }}, надбавка.', {'Премия': ''})
    assert text == 'Оклад, надбавка.'


def test_empty_first_value_trims_leading_space(tmp_path: Path) -> None:
    text = _render_body(tmp_path, '{{ Звание }} Петров назначен.', {'Звание': ''})
    assert text == 'Петров назначен.'


def test_nonempty_value_keeps_terminator(tmp_path: Path) -> None:
    text = _render_body(tmp_path, 'Звание: {{ Звание }}.', {'Звание': 'майор'})
    assert text == 'Звание: майор.'


def test_legitimate_terminator_not_removed(tmp_path: Path) -> None:
    # Нет пустых подстановок — точка в конце предложения не трогается.
    text = _render_body(tmp_path, 'Иванов окончил институт.', {})
    assert text == 'Иванов окончил институт.'


def test_intentional_spaces_preserved_without_empty(tmp_path: Path) -> None:
    # T026 issue 4: намеренная последовательность пробелов (подпись) без пустых
    # подстановок — НЕ трогается (программа не самодействует вне подстановок).
    text = _render_body(tmp_path, 'Подпись:          И.И. Петров', {})
    assert text == 'Подпись:          И.И. Петров'


def test_intentional_spaces_preserved_with_empty_elsewhere(tmp_path: Path) -> None:
    # Пустая подстановка чистит только СВОЙ дублированный пробел; намеренные
    # пробелы в другом месте абзаца остаются нетронутыми.
    text = _render_body(
        tmp_path, '{{ Звание }} Иванов     —     директор', {'Звание': ''}
    )
    assert text == 'Иванов     —     директор'


def test_empty_value_does_not_touch_dash(tmp_path: Path) -> None:
    # Тире (не закрывающая пунктуация) не затрагивается; ведущий пробел убран.
    text = _render_body(tmp_path, '{{ Звание }} — Петров', {'Звание': ''})
    assert text == '— Петров'


def test_filename_strips_empty_marker(tmp_path: Path) -> None:
    # Маркер пустой подстановки не должен попасть в имя файла.
    name = render_filename(
        'Приказ_{{ Звание }}.docx', build_context(Person(cells={'Звание': ''}))
    )
    assert EMPTY_MARKER not in name
    assert name == 'Приказ_.docx'


def test_filename_keeps_html_special_chars_literal() -> None:
    # T011: общий с телом env держит autoescape=True, но в имени файла
    # `&`/`'` должны остаться буквальными, а не уехать в `&amp;`/`&#39;`.
    name = render_filename(
        'Приказ_{{ Орг }}.docx', build_context(Person(cells={'Орг': "Рога & Ко 'A'"}))
    )
    assert name == "Приказ_Рога & Ко 'A'.docx"


def test_body_keeps_html_escaping_for_special_chars(tmp_path: Path) -> None:
    # T011-страховка: фикс имени файла не должен ослабить экранирование тела.
    # Значение с `&` обязано сохраниться как валидный XML (`&amp;`), иначе docx
    # испортится; python-docx читает его обратно как `&`.
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('Организация: {{ Орг }}')
    doc.save(template)

    out = tmp_path / 'out.docx'
    render_document(template, build_context(Person(cells={'Орг': 'Рога & Ко'})), out)

    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'Организация: Рога & Ко' in text
