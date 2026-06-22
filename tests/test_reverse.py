"""Тесты обратной генерации шаблона (`dyak reverse`, T007 фазы 1–2)."""

from __future__ import annotations

import csv
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from typer.testing import CliRunner

from dyak.cli import app, reverse_template
from dyak.columns import (
    KEY_NAME,
    KEY_PATRONYMIC,
    KEY_SURNAME,
    NAME,
    PATRONYMIC,
    POSITION,
    SURNAME,
)
from dyak.domain import Person
from dyak.inflection import PetrovichInflector, PymorphyInflector
from dyak.render.context import build_context
from dyak.render.engine import render_document
from dyak.reverse import FindingKind, build_template, format_report
from dyak.reverse.candidates import build_candidates
from dyak.reverse.docx_rewrite import flatten_runs, rewrite_paragraph
from dyak.reverse.matcher import find_spans
from dyak.reverse.report import Finding, ReverseReport

_FIXTURE = Path(__file__).parent / 'fixtures' / 'declension.csv'
_CASE_RUS = ('им', 'рд', 'дт', 'вн', 'тв', 'пр')

# Роли для отдельных колонок Фамилия/Имя/Отчество (scenario 1, по частям).
_PART_ROLES = {SURNAME: KEY_SURNAME, NAME: KEY_NAME, PATRONYMIC: KEY_PATRONYMIC}


def _person(**cells: str) -> Person:
    return Person(cells=dict(cells))


def _paragraph_text(doc: Document) -> str:
    return '\n'.join(p.text for p in doc.paragraphs)


@pytest.fixture(scope='module')
def fio_inflector() -> PetrovichInflector:
    return PetrovichInflector()


@pytest.fixture(scope='module')
def pos_inflector() -> PymorphyInflector:
    return PymorphyInflector()


# --- candidates ---------------------------------------------------------------


def test_build_candidates_skips_empty_cells() -> None:
    person = _person(Фамилия='Иванов', Отчество='', Должность='  ')
    candidates = build_candidates(person)
    keys = {c.key for c in candidates}
    assert keys == {'Фамилия'}
    (cand,) = candidates
    (form,) = cand.forms
    assert form.text == 'Иванов'
    assert form.tag == '{{ Фамилия }}'
    assert not form.ambiguous


# --- matcher ------------------------------------------------------------------


def test_find_spans_exact_match() -> None:
    candidates = build_candidates(_person(Фамилия='Иванов'))
    (match,) = find_spans('Назначить Иванов', candidates)
    assert (match.start, match.end) == (10, 16)
    assert match.form.tag == '{{ Фамилия }}'


def test_find_spans_respects_word_boundary() -> None:
    # Значение «Иван» не должно совпасть внутри «Иванов» (склонённые/составные
    # формы — забота фазы 2, не точного матча).
    candidates = build_candidates(_person(Имя='Иван'))
    assert find_spans('Иванов', candidates) == []


def test_find_spans_inflected_form_not_matched() -> None:
    # «Иванову» (дат. падеж) не совпадает с именительным «Иванов».
    candidates = build_candidates(_person(Фамилия='Иванов'))
    assert find_spans('вручить Иванову', candidates) == []


def test_find_spans_prefers_longer_on_overlap() -> None:
    person = _person(Должность='главный бухгалтер', Слово='бухгалтер')
    matches = find_spans('главный бухгалтер', build_candidates(person))
    assert len(matches) == 1
    assert matches[0].candidate.key == 'Должность'


def test_find_spans_multiple_occurrences() -> None:
    candidates = build_candidates(_person(Город='Москва'))
    matches = find_spans('Москва — город Москва', candidates)
    assert [m.start for m in matches] == [0, 15]


# --- docx run-split rewriting -------------------------------------------------


def test_rewrite_run_split_keeps_tag_intact(tmp_path: Path) -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run('Назначить Ива')  # значение «Иванов» разрезано по run'ам
    para.add_run('нов на должность')

    full, bounds = flatten_runs(para)
    matches = find_spans(full, build_candidates(_person(Фамилия='Иванов')))
    rewrite_paragraph(para, full, bounds, matches)

    # Тег целиком в одном run — docxtpl увидит его неразрезанным.
    assert any('{{ Фамилия }}' in run.text for run in para.runs)
    assert para.text == 'Назначить {{ Фамилия }} на должность'

    # И forward-рендер шаблона воспроизводит исходный текст.
    template = tmp_path / 'tpl.docx'
    doc.save(template)
    out = tmp_path / 'out.docx'
    render_document(template, build_context(_person(Фамилия='Иванов')), out)
    assert 'Назначить Иванов на должность' in _paragraph_text(Document(out))


# --- engine end-to-end --------------------------------------------------------


def _sample_doc(path: Path, text: str) -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_build_template_replaces_flat_values_roundtrip(tmp_path: Path) -> None:
    person = _person(
        Фамилия='Иванов', Имя='Пётр', Должность='директор', Дата_начала='01.07.2026'
    )
    sample = _sample_doc(
        tmp_path / 'sample.docx',
        'Назначить Иванов Пётр на должность директор с 01.07.2026.',
    )
    document, report = build_template(sample, person)

    body = _paragraph_text(document)
    assert '{{ Фамилия }}' in body
    assert '{{ Дата_начала }}' in body
    assert report.replaced_count == 4
    assert report.of_kind(FindingKind.NOT_FOUND) == []

    # Round-trip: собранный шаблон воспроизводит исходный документ.
    template = tmp_path / 'tpl.docx'
    document.save(template)
    out = tmp_path / 'out.docx'
    render_document(template, build_context(person), out)
    assert (
        'Назначить Иванов Пётр на должность директор с 01.07.2026.'
        in _paragraph_text(Document(out))
    )


def test_build_template_reports_not_found(tmp_path: Path) -> None:
    person = _person(Фамилия='Иванов', Отчество='Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Уведомить Иванов.')
    _document, report = build_template(sample, person)

    not_found = report.of_kind(FindingKind.NOT_FOUND)
    assert any('Семёнович' in f.message for f in not_found)


def test_build_template_reports_unmatched_data(tmp_path: Path) -> None:
    # Дата в документе, которой нет в строке → unmatched_text.
    person = _person(Фамилия='Иванов')
    sample = _sample_doc(
        tmp_path / 'sample.docx', 'Иванов, приказ от 09.09.2099.'
    )
    _document, report = build_template(sample, person)

    unmatched = report.of_kind(FindingKind.UNMATCHED_TEXT)
    assert any('09.09.2099' in f.message for f in unmatched)


def test_build_template_rewrites_table_cells(tmp_path: Path) -> None:
    # Значения во вложенной таблице тоже переписываются.
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run('директор Иванов')
    sample = tmp_path / 'sample.docx'
    doc.save(sample)

    person = _person(Фамилия='Иванов', Должность='директор')
    document, _report = build_template(sample, person)
    cell_text = document.tables[0].cell(0, 0).text
    assert '{{ Фамилия }}' in cell_text
    assert '{{ Должность }}' in cell_text


# --- report formatting --------------------------------------------------------


def test_format_report_groups_sections() -> None:
    report = ReverseReport(
        findings=[
            Finding(FindingKind.REPLACED, 'Фамилия: «Иванов» → {{ Фамилия }}'),
            Finding(FindingKind.NOT_FOUND, 'Отчество: «Семёнович» не найдено'),
        ]
    )
    text = format_report(report)
    assert 'Заменено значений: 1' in text
    assert 'Заменено:' in text
    assert 'Не найдено в документе:' in text


# --- phase 2: decline-and-match -----------------------------------------------


def _fixture_rows() -> list[dict[str, str]]:
    with _FIXTURE.open(encoding='utf-8') as fh:
        return list(csv.DictReader(fh))


def _parts_template(
    sample: Path,
    person: Person,
    fio: PetrovichInflector,
) -> tuple[Document, ReverseReport]:
    """Шаблон в режиме отдельных колонок ФИО (scenario 1, по частям)."""
    return build_template(sample, person, roles=_PART_ROLES, inflector=fio)


def _parts_context(person: Person, fio: PetrovichInflector) -> dict[str, object]:
    return build_context(person, roles=_PART_ROLES, inflector=fio)


def test_declension_parts_oblique_case_roundtrip(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(
        tmp_path / 'sample.docx', 'Вручить Иванову Петру Семёновичу уведомление.'
    )
    document, report = _parts_template(sample, person, fio_inflector)

    body = _paragraph_text(document)
    assert '{{ Фамилия | дт }}' in body
    assert '{{ Имя | дт }}' in body
    assert '{{ Отчество | дт }}' in body
    assert report.of_kind(FindingKind.NOT_FOUND) == []

    template = tmp_path / 'tpl.docx'
    document.save(template)
    out = tmp_path / 'out.docx'
    render_document(template, _parts_context(person, fio_inflector), out)
    assert (
        'Вручить Иванову Петру Семёновичу уведомление.'
        in _paragraph_text(Document(out))
    )


def test_declension_nominative_has_no_filter(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов Пётр Семёнович — директор.')
    document, _report = _parts_template(sample, person, fio_inflector)
    body = _paragraph_text(document)
    assert '{{ Фамилия }}' in body
    assert '| ' not in body  # именительный падеж — теги без фильтра


def test_homonymous_case_marked_ambiguous(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    # Муж. фамилия «Иванова» = рд = вн → берём первый (рд) и помечаем ambiguous.
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Касательно Иванова Петра.')
    document, report = _parts_template(sample, person, fio_inflector)

    body = _paragraph_text(document)
    assert '{{ Фамилия | рд }}' in body
    ambiguous = report.of_kind(FindingKind.AMBIGUOUS)
    assert any('Фамилия' in f.message for f in ambiguous)


def test_ambiguous_run_is_highlighted(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Касательно Иванова.')
    document, _report = _parts_template(sample, person, fio_inflector)
    highlighted = [
        run
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.font.highlight_color is not None
    ]
    assert highlighted, 'неоднозначный падеж должен быть подсвечен'


@pytest.mark.parametrize('row', _fixture_rows(), ids=lambda r: f"{r['surname']}-{r['gender']}")
@pytest.mark.parametrize('rus', _CASE_RUS)
def test_fullname_column_roundtrip_all_cases(
    row: dict[str, str],
    rus: str,
    tmp_path: Path,
    fio_inflector: PetrovichInflector,
) -> None:
    """Регрессия: целое ФИО в каждом падеже даёт шаблон, воспроизводящий образец."""
    fullname_nomn = row['им']
    target = row[rus]
    person = _person(ФИО=fullname_nomn)
    sample = _sample_doc(tmp_path / 'sample.docx', f'Документ для {target}.')
    document, _report = build_template(
        sample,
        person,
        fullname_source='ФИО',
        roles=_PART_ROLES,
        inflector=fio_inflector,
    )
    assert '{{ ФИО' in _paragraph_text(document)

    template = tmp_path / 'tpl.docx'
    document.save(template)
    out = tmp_path / 'out.docx'
    context = build_context(
        person, fullname_source='ФИО', roles=_PART_ROLES, inflector=fio_inflector
    )
    render_document(template, context, out)
    assert f'Документ для {target}.' in _paragraph_text(Document(out))


def test_position_oblique_case_roundtrip(
    tmp_path: Path, fio_inflector: PetrovichInflector, pos_inflector: PymorphyInflector
) -> None:
    person = _person(Должность='директор')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Назначить директором.')
    document, _report = build_template(
        sample,
        person,
        roles={POSITION: 'Должность'},
        inflector=fio_inflector,
        position_inflector=pos_inflector,
    )
    assert '{{ Должность | тв }}' in _paragraph_text(document)

    template = tmp_path / 'tpl.docx'
    document.save(template)
    out = tmp_path / 'out.docx'
    context = build_context(
        person,
        roles={POSITION: 'Должность'},
        inflector=fio_inflector,
        position_inflector=pos_inflector,
    )
    render_document(template, context, out)
    assert 'Назначить директором.' in _paragraph_text(Document(out))


def test_initials_recognized_best_effort(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Подпись: Иванов П. С.')
    document, _report = _parts_template(sample, person, fio_inflector)
    body = _paragraph_text(document)
    assert '{{ ФИО.инициалы }}' in body

    template = tmp_path / 'tpl.docx'
    document.save(template)
    out = tmp_path / 'out.docx'
    render_document(template, _parts_context(person, fio_inflector), out)
    assert 'Подпись: Иванов П. С.' in _paragraph_text(Document(out))


def test_fullname_prefers_whole_over_parts(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    # Колонка «ФИО» + полное имя подряд → единый тег, не три части (Q3).
    person = _person(ФИО='Иванов Пётр Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов Пётр Семёнович назначен.')
    document, _report = build_template(
        sample,
        person,
        fullname_source='ФИО',
        roles=_PART_ROLES,
        inflector=fio_inflector,
    )
    body = _paragraph_text(document)
    assert '{{ ФИО }}' in body
    assert '{{ Фамилия }}' not in body


def test_secondary_parts_do_not_report_not_found(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    # При колонке «ФИО» производные части — вторичные: не шумят not_found.
    person = _person(ФИО='Иванов Пётр Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов Пётр Семёнович назначен.')
    _document, report = build_template(
        sample,
        person,
        fullname_source='ФИО',
        roles=_PART_ROLES,
        inflector=fio_inflector,
    )
    not_found = report.of_kind(FindingKind.NOT_FOUND)
    assert not any('Фамилия' in f.message for f in not_found)


def test_surname_alone_tagged_with_fullname_column(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    # Отдельное упоминание фамилии при колонке «ФИО» всё равно тегируется
    # (иначе при перегенерации осталась бы фамилия исходного человека).
    person = _person(ФИО='Иванов Пётр Семёнович')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Уведомить Иванова о решении.')
    document, _report = build_template(
        sample,
        person,
        fullname_source='ФИО',
        roles=_PART_ROLES,
        inflector=fio_inflector,
    )
    assert '{{ Фамилия | рд }}' in _paragraph_text(document)


# --- phase 3: round-trip verify -----------------------------------------------


def test_roundtrip_clean_build_has_no_mismatch(
    tmp_path: Path, fio_inflector: PetrovichInflector
) -> None:
    # Корректно собранный шаблон воспроизводит образец — расхождений нет.
    person = _person(Фамилия='Иванов', Имя='Пётр', Отчество='Семёнович')
    sample = _sample_doc(
        tmp_path / 'sample.docx', 'Вручить Иванову Петру Семёновичу уведомление.'
    )
    _document, report = _parts_template(sample, person, fio_inflector)
    assert report.of_kind(FindingKind.ROUNDTRIP_MISMATCH) == []


def test_roundtrip_detects_stray_template_syntax(tmp_path: Path) -> None:
    # Образец с {% raw %}-шпаргалкой: forward-рендер срежет управляющие теги,
    # текст разойдётся с оригиналом → roundtrip_mismatch (шаблон всё равно собран).
    person = _person(Фамилия='Иванов')
    sample = _sample_doc(
        tmp_path / 'sample.docx',
        'Памятка {% raw %}{{ Дата }}{% endraw %} для Иванов.',
    )
    document, report = build_template(sample, person)
    assert '{{ Фамилия }}' in _paragraph_text(document)  # шаблон собран
    mismatches = report.of_kind(FindingKind.ROUNDTRIP_MISMATCH)
    assert mismatches
    assert 'ожидалось' in mismatches[0].message


def test_roundtrip_render_failure_recorded_not_raised(tmp_path: Path) -> None:
    # Забытый неизвестный тег в образце: forward-рендер падает StrictUndefined,
    # но reverse best-effort — фиксируем находкой, шаблон всё равно сохранён.
    person = _person(Фамилия='Иванов')
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов, см. {{ Неизвестно }}.')
    document, report = build_template(sample, person)
    assert '{{ Фамилия }}' in _paragraph_text(document)
    mismatches = report.of_kind(FindingKind.ROUNDTRIP_MISMATCH)
    assert any('не выполнена' in f.message for f in mismatches)


# --- CLI ----------------------------------------------------------------------


def _xlsx(path: Path, headers: list[str], rows: list[list[str]]) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)
    return path


def test_cli_reverse_builds_template(tmp_path: Path) -> None:
    table = _xlsx(
        tmp_path / 'emp.xlsx',
        ['Фамилия', 'Имя', 'Должность'],
        [['Иванов', 'Пётр', 'директор']],
    )
    sample = _sample_doc(tmp_path / 'sample.docx', 'Назначить Иванов на директор.')
    out = tmp_path / 'tpl.docx'
    result = CliRunner().invoke(
        app,
        [
            'reverse',
            '--doc', str(sample),
            '--table', str(table),
            '--row', '1',
            '--out', str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    body = _paragraph_text(Document(out))
    assert '{{ Фамилия }}' in body
    assert '{{ Должность }}' in body


def test_cli_reverse_row_out_of_range_exits_one(tmp_path: Path) -> None:
    table = _xlsx(tmp_path / 'emp.xlsx', ['Фамилия'], [['Иванов']])
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов.')
    result = CliRunner().invoke(
        app,
        [
            'reverse',
            '--doc', str(sample),
            '--table', str(table),
            '--row', '5',
            '--out', str(tmp_path / 'tpl.docx'),
        ],
    )
    assert result.exit_code == 1
    assert 'диапазона' in result.output


def test_reverse_template_helper_returns_report(tmp_path: Path) -> None:
    table = _xlsx(tmp_path / 'emp.xlsx', ['Фамилия'], [['Иванов']])
    sample = _sample_doc(tmp_path / 'sample.docx', 'Иванов.')
    report = reverse_template(
        sample, table, tmp_path / 'tpl.docx', None, None, 1
    )
    assert report.replaced_count == 1
