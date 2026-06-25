"""
Регрессия «склонение по умолчанию» (T024): сквозной прогон приказа.

Фикстура — реальный пример (`примеры для дьяка`): приказ с колонками
Подразделение/Должность/Звание/ФИО и шаблоном, бьющим по всем движкам сразу
(`{{ Должность|дт }} {{ Подразделение|вн }} {{ Звание|дт }} {{ Фамилия|дт }}`).
Покрывает: generic-склонение любой текстовой колонки, растворение T020
(«Должность» склоняется даже когда «Подразделение» забрало роль position),
замороженный родительный хвост, одушевлённость винительного, составную
должность через дефис, ФИО через petrovich+пол, звание спец-движком.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

import openpyxl
import pytest
from docx import Document

from dyak.cli import generate_documents
from dyak.domain import Case, Person
from dyak.inflection import Phrase, PhraseInflector
from dyak.render.context import build_context

if TYPE_CHECKING:
    from pathlib import Path

_HEADERS = ['№ п/п', 'Подразделение', 'Должность', 'Звание', 'ФИО', 'л/н']
_ROWS = [
    (1, '1 мотострелковый батальон', 'командир', 'подполковник',
     'ПУПКИН Василий Федорович', 'ХУ-123456'),
    (2, 'рота связи', 'телефонист', 'ефрейтор',
     'БИДЕ Абдуразах Аббасович', 'ЧС-987654'),
    (3, 'взвод охраны', 'стрелок', 'рядовой',
     'СЕМЕНОВ Петр Семенович', 'КУ-864532'),
    (4, 'танковый батальон', 'механик-водитель', 'младший сержант',
     'БИВЕНЬ Абрам Моисеевич', 'ПР-109283'),
]  # fmt: skip

_TEMPLATE = '{{ Должность|дт }} {{ Подразделение|вн }} {{ Звание|дт }} {{ Фамилия|дт }}'

# Эталон по примеру приказа (грамматически корректный текст по всем строкам).
_EXPECTED = {
    'ПУПКИН': 'командиру 1 мотострелковый батальон подполковнику ПУПКИНУ',
    'БИДЕ': 'телефонисту роту связи ефрейтору БИДЕ',  # БИДЕ несклоняемая
    'СЕМЕНОВ': 'стрелку взвод охраны рядовому СЕМЕНОВУ',
    # БИВЕНЬ — фамилия-нарицательное (pymorphy не знает её как фамилию): по
    # умолчанию остаётся в именительном (T027), как в офиц. документах.
    'БИВЕНЬ': 'механику-водителю танковый батальон младшему сержанту БИВЕНЬ',
}


@pytest.fixture
def order_fixture(tmp_path: Path) -> tuple[Path, Path]:
    """Записать таблицу-приказ и шаблон; вернуть их пути."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(_HEADERS)
    for row in _ROWS:
        sheet.append(list(row))
    table = tmp_path / 'order.xlsx'
    workbook.save(str(table))

    document = Document()
    document.add_paragraph(_TEMPLATE)
    template = tmp_path / 'order.docx'
    document.save(str(template))
    return table, template


def test_order_declension_matches_reference(
    order_fixture: tuple[Path, Path], tmp_path: Path
) -> None:
    """Каждая из 4 строк приказа даёт грамматически корректный текст."""
    table, template = order_fixture
    out = tmp_path / 'out'
    paths = generate_documents(table, template, out, None, None, '{{ Фамилия }}.docx')
    by_surname = {
        path.stem: ' '.join(
            par.text for par in Document(str(path)).paragraphs if par.text.strip()
        )
        for path in paths
    }
    assert by_surname == _EXPECTED


def test_unrole_text_column_declines_by_default() -> None:
    """Колонка без объявленной роли склоняется generic-движком (T024)."""
    person = Person(cells={'Основание': 'служебная записка'})
    context = build_context(person, position_inflector=PhraseInflector())
    value = context['Основание']
    assert isinstance(value, Phrase)
    assert value.inflect(Case.GENT) == 'служебной записки'


def test_literal_value_is_not_wrapped() -> None:
    """Коды/личные номера/даты остаются литералом (не оборачиваются)."""
    person = Person(cells={'л/н': 'ХУ-123456', 'Дата': '12.05.2020', 'Должность': 'стрелок'})
    context = build_context(person, position_inflector=PhraseInflector())
    assert context['л/н'] == 'ХУ-123456'  # литерал — строка
    assert context['Дата'] == '12.05.2020'  # дата — строка
    assert isinstance(context['Должность'], Phrase)  # текст — обёрнут
