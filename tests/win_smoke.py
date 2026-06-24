"""
Windows-smoke ядра: проверяет КОРРЕКТНОСТЬ склонения собранным exe (T021).

Retrospective 0.3.0: CI ассертил только факт создания файлов, а не их
содержимое — через эту дыру в релиз ушёл баг T021 (на русской Windows
`petrovich` читал `rules.json` в `cp1251` и фамилия не склонялась). Этот
скрипт закрывает дыру: строит мини-фикстуру, даёт собранному exe её
сгенерировать и ассертит, что в готовом документе фамилия и должность
просклонены в дательный падеж.

Используется из `build-windows.yml` в трёх шагах:

    uv run python tests/win_smoke.py make   <work_dir>
    dist/dyak-portable.exe generate --table … --template … --out … *> log
    uv run python tests/win_smoke.py check  <out_dir>

`make` пишет `table.xlsx` + `template.docx`; `check` падает с кодом 1, если
ожидаемых склонённых форм нет в выходных документах. Запускается и на Linux
(из исходников) — для разработки и юнит-проверки самой логики.
"""

from __future__ import annotations

import sys
from pathlib import Path

import openpyxl
from docx import Document

# Одна строка-фикстура: колонка «ФИО» (petrovich) + «Должность» (фраз-движок).
_HEADERS = ('ФИО', 'Должность')
_ROW = ('ПУПКИН Василий Федорович', 'командир')

# Шаблон бьёт по обоим движкам склонения в дательном падеже.
_TEMPLATE_TEXT = 'Объявить выговор {{ Должность|дт }} {{ Фамилия|дт }}.'

# Формы, которые ОБЯЗАНЫ появиться в готовом документе (дательный падеж).
# «ПУПКИНУ» — основной таргет T021 (petrovich); «командиру» — страховка, что
# фраз-движок жив. Несклонённые «ПУПКИН»/«командир» → smoke красный.
_EXPECTED = ('ПУПКИНУ', 'командиру')


def build_fixtures(work_dir: Path) -> tuple[Path, Path]:
    """Записать `table.xlsx` и `template.docx` в `work_dir`; вернуть их пути."""
    work_dir.mkdir(parents=True, exist_ok=True)
    table = work_dir / 'table.xlsx'
    template = work_dir / 'template.docx'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(list(_HEADERS))
    sheet.append(list(_ROW))
    workbook.save(str(table))

    document = Document()
    document.add_paragraph(_TEMPLATE_TEXT)
    document.save(str(template))
    return table, template


def find_problems(out_dir: Path) -> list[str]:
    """Вернуть список ненайденных ожидаемых форм (пусто = склонение корректно)."""
    docs = sorted(out_dir.glob('*.docx'))
    if not docs:
        return [f'в {out_dir} нет ни одного .docx — генерация не отработала']
    text = '\n'.join(
        paragraph.text for doc in docs for paragraph in Document(str(doc)).paragraphs
    )
    return [form for form in _EXPECTED if form not in text]


def _force_utf8_output() -> None:
    """UTF-8 на собственный stdout/stderr — иначе кириллица падает на cp1252 CI."""
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, 'reconfigure', None)
        if callable(reconfigure):
            reconfigure(encoding='utf-8', errors='replace')


_ARGC = 2  # ровно команда + путь


def main(argv: list[str]) -> int:
    """Диспетчер: `make <dir>` строит фикстуру, `check <dir>` проверяет вывод."""
    _force_utf8_output()
    if len(argv) != _ARGC:
        sys.stderr.write('usage: win_smoke.py make|check <dir>\n')
        return 2
    command, raw_path = argv
    path = Path(raw_path)
    if command == 'make':
        table, template = build_fixtures(path)
        sys.stdout.write(f'fixtures: {table}, {template}\n')
        return 0
    if command == 'check':
        problems = find_problems(path)
        if problems:
            sys.stderr.write('SMOKE FAILED — склонение неверно:\n')
            for problem in problems:
                sys.stderr.write(f'  - не найдено: {problem}\n')
            return 1
        sys.stdout.write('SMOKE OK: фамилия и должность просклонены\n')
        return 0
    sys.stderr.write(f'unknown command: {command}\n')
    return 2


if __name__ == '__main__':
    raise SystemExit(main(sys.argv[1:]))
