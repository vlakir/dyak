"""Тесты PDF-экспорта через LibreOffice (T012): поиск soffice + конвертация."""

from __future__ import annotations

import subprocess
from pathlib import Path
from typing import TYPE_CHECKING

import pytest
from docx import Document
from openpyxl import Workbook
from typer.testing import CliRunner

from dyak import pdf
from dyak.cli import app
from dyak.errors import PdfExportError
from dyak.pdf import export_to_pdf, find_soffice

if TYPE_CHECKING:
    from collections.abc import Sequence

runner = CliRunner()

_HAS_SOFFICE = find_soffice() is not None
_needs_soffice = pytest.mark.skipif(not _HAS_SOFFICE, reason='LibreOffice не установлен')


def _fake_run(
    *,
    returncode: int = 0,
    stderr: str = '',
    make_pdfs: Sequence[Path] = (),
) -> object:
    """Фабрика подмены `subprocess.run`: создаёт PDF-файлы и возвращает код."""

    def run(command: list[str], **_: object) -> subprocess.CompletedProcess[str]:
        for target in make_pdfs:
            target.write_bytes(b'%PDF-1.7 fake')
        return subprocess.CompletedProcess(command, returncode, '', stderr)

    return run


def test_find_soffice_in_path(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        pdf.shutil,
        'which',
        lambda name: '/usr/bin/soffice' if name == 'soffice' else None,
    )
    assert find_soffice() == Path('/usr/bin/soffice')


def test_find_soffice_none_when_absent(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(pdf.shutil, 'which', lambda _: None)
    monkeypatch.setattr(pdf, '_candidate_paths', list)
    assert find_soffice() is None


def test_find_soffice_falls_back_to_candidate(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    binary = tmp_path / 'soffice'
    binary.write_text('', encoding='utf-8')
    monkeypatch.setattr(pdf.shutil, 'which', lambda _: None)
    monkeypatch.setattr(
        pdf,
        '_candidate_paths',
        lambda: [tmp_path / 'missing', binary],
    )
    assert find_soffice() == binary


def test_candidate_paths_include_windows(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv('ProgramFiles', r'C:\Program Files')
    paths = [str(path) for path in pdf._candidate_paths()]
    assert any(p.endswith('soffice.exe') and 'LibreOffice' in p for p in paths)


def test_export_empty_returns_empty(tmp_path: Path) -> None:
    assert export_to_pdf([], tmp_path) == []


def test_export_raises_when_soffice_missing(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(pdf, 'find_soffice', lambda: None)
    with pytest.raises(PdfExportError, match='не найден'):
        export_to_pdf([tmp_path / 'a.docx'], tmp_path)


def test_export_success_mocked(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx = [tmp_path / 'a.docx', tmp_path / 'b.docx']
    monkeypatch.setattr(pdf, 'find_soffice', lambda: Path('/usr/bin/soffice'))
    monkeypatch.setattr(
        pdf.subprocess,
        'run',
        _fake_run(make_pdfs=[tmp_path / 'a.pdf', tmp_path / 'b.pdf']),
    )
    result = export_to_pdf(docx, tmp_path)
    assert [path.name for path in result] == ['a.pdf', 'b.pdf']


def test_export_raises_on_nonzero(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(pdf, 'find_soffice', lambda: Path('/usr/bin/soffice'))
    monkeypatch.setattr(
        pdf.subprocess,
        'run',
        _fake_run(returncode=1, stderr='boom'),
    )
    with pytest.raises(PdfExportError, match='код 1'):
        export_to_pdf([tmp_path / 'a.docx'], tmp_path)


def test_export_raises_when_pdf_missing(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # soffice «успешен», но файла нет — ловим явно, а не молча.
    monkeypatch.setattr(pdf, 'find_soffice', lambda: Path('/usr/bin/soffice'))
    monkeypatch.setattr(pdf.subprocess, 'run', _fake_run())
    with pytest.raises(PdfExportError, match='не появились'):
        export_to_pdf([tmp_path / 'a.docx'], tmp_path)


def _make_inputs(tmp_path: Path) -> tuple[Path, Path]:
    table = tmp_path / 'table.xlsx'
    wb = Workbook()
    ws = wb.active
    ws.append(['Фамилия', 'Имя'])
    ws.append(['Иванов', 'Пётр'])
    wb.save(table)
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('Сотрудник {{ Фамилия }} {{ Имя }}.')
    doc.save(template)
    return table, template


def test_cli_pdf_missing_soffice_degrades(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # --pdf без LibreOffice → понятная ошибка и код 1, docx всё равно создан.
    monkeypatch.setattr(pdf, 'find_soffice', lambda: None)
    table, template = _make_inputs(tmp_path)
    out = tmp_path / 'out'
    result = runner.invoke(
        app,
        ['generate', '--table', str(table), '--template', str(template),
         '--out', str(out), '--pdf'],
    )
    assert result.exit_code == 1
    assert 'LibreOffice' in result.stderr
    assert list(out.glob('*.docx'))  # docx сгенерированы до сбоя PDF


@_needs_soffice
def test_export_real_pdf(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph('Привет, мир.')
    source = tmp_path / 'doc.docx'
    doc.save(source)

    result = export_to_pdf([source], tmp_path)

    assert len(result) == 1
    assert result[0].exists()
    assert result[0].read_bytes().startswith(b'%PDF')


@_needs_soffice
def test_cli_generate_pdf(tmp_path: Path) -> None:
    table, template = _make_inputs(tmp_path)
    out = tmp_path / 'out'
    result = runner.invoke(
        app,
        ['generate', '--table', str(table), '--template', str(template),
         '--out', str(out), '--pdf'],
    )
    assert result.exit_code == 0
    assert 'PDF' in result.stdout
    assert list(out.glob('*.pdf'))
