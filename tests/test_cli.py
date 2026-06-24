"""Тесты команды generate (T006: подстановка по заголовкам, --filename)."""

from __future__ import annotations

import logging
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from typer.testing import CliRunner

from dyak.cli import _gender_overrides, app, configure_stdio, generate_documents
from dyak.config import Config
from dyak.domain import Gender

_HEADERS = ['Фамилия', 'Имя', 'Отчество', 'Должность', 'Дата начала']
_ROWS = [
    ['Иванов', 'Пётр', 'Семёнович', 'директор', '01.07.2026'],
    ['Петрова', 'Анна', 'Сергеевна', 'главный бухгалтер', '02.07.2026'],
]


def _make_xlsx(path: Path, rows: list[list[str]], headers: list[str]) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)
    return path


def _make_template(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph(
        'Назначить {{ Фамилия }} {{ Имя }} на должность '
        '{{ Должность }} с {{ Дата_начала }}.',
    )
    doc.save(path)
    return path


def _fixtures(
    tmp_path: Path,
    rows: list[list[str]] | None = None,
    headers: list[str] | None = None,
) -> tuple[Path, Path]:
    table = _make_xlsx(
        tmp_path / 'emp.xlsx',
        rows if rows is not None else _ROWS,
        headers if headers is not None else _HEADERS,
    )
    template = _make_template(tmp_path / 'tpl.docx')
    return table, template


def test_generate_with_filename_template(tmp_path: Path) -> None:
    table, template = _fixtures(tmp_path)
    out = tmp_path / 'out'
    written = generate_documents(
        table,
        template,
        out,
        config=None,
        sheet=None,
        filename='Приказ_{{ Фамилия }}.docx',
    )
    assert sorted(p.name for p in written) == [
        'Приказ_Иванов.docx',
        'Приказ_Петрова.docx',
    ]
    text = '\n'.join(p.text for p in Document(out / 'Приказ_Иванов.docx').paragraphs)
    assert 'Назначить Иванов Пётр на должность директор с 01.07.2026.' in text


def test_generate_default_filename_from_fio(tmp_path: Path) -> None:
    table, template = _fixtures(tmp_path)
    out = tmp_path / 'out'
    written = generate_documents(
        table, template, out, config=None, sheet=None, filename=None,
    )
    assert sorted(p.name for p in written) == [
        'Иванов_Пётр_Семёнович.docx',
        'Петрова_Анна_Сергеевна.docx',
    ]


def test_generate_ordinal_fallback_without_roles(tmp_path: Path) -> None:
    headers = ['Дата начала', 'Номер приказа']
    rows = [['01.07.2026', '17'], ['02.07.2026', '18']]
    table = _make_xlsx(tmp_path / 'emp.xlsx', rows, headers)
    template = tmp_path / 'tpl.docx'
    Document().save(template)
    out = tmp_path / 'out'
    written = generate_documents(
        table, template, out, config=None, sheet=None, filename=None,
    )
    assert sorted(p.name for p in written) == ['Документ_1.docx', 'Документ_2.docx']


def test_generate_from_single_fullname_column(tmp_path: Path) -> None:
    headers = ['ФИО', 'Должность']
    rows = [['Иванов Пётр Семёнович', 'директор'], ['Петрова Анна', 'бухгалтер']]
    table = _make_xlsx(tmp_path / 'emp.xlsx', rows, headers)
    template = tmp_path / 'tpl.docx'
    doc = Document()
    doc.add_paragraph('{{ ФИО }} — {{ Фамилия }} {{ Имя }} {{ Отчество }}')
    doc.save(template)
    out = tmp_path / 'out'
    written = generate_documents(
        table, template, out, config=None, sheet=None, filename=None,
    )
    # Дефолтное имя файла собирается из разобранных частей ФИО; у строки из
    # двух слов отчество пустое и хвостового подчёркивания нет.
    assert sorted(p.name for p in written) == [
        'Иванов_Пётр_Семёнович.docx',
        'Петрова_Анна.docx',
    ]
    text = '\n'.join(
        p.text for p in Document(out / 'Иванов_Пётр_Семёнович.docx').paragraphs
    )
    assert 'Иванов Пётр Семёнович — Иванов Пётр Семёнович' in text


def test_generate_creates_out_dir(tmp_path: Path) -> None:
    table, template = _fixtures(tmp_path)
    out = tmp_path / 'nested' / 'out'
    generate_documents(
        table, template, out, config=None, sheet=None, filename='{{ Фамилия }}.docx',
    )
    assert out.is_dir()


def test_generate_resolves_name_collision(tmp_path: Path) -> None:
    rows = [
        ['Иванов', 'Пётр', 'Семёнович', 'директор', '01.07.2026'],
        ['Иванов', 'Иван', 'Иванович', 'инженер', '02.07.2026'],
    ]
    table, template = _fixtures(tmp_path, rows)
    out = tmp_path / 'out'
    written = generate_documents(
        table, template, out, config=None, sheet=None, filename='{{ Фамилия }}.docx',
    )
    assert sorted(p.name for p in written) == ['Иванов.docx', 'Иванов_2.docx']


def test_module_entrypoint_exposes_app() -> None:
    import dyak.__main__ as entry

    assert entry.app is not None


def test_cli_generate_exit_zero_zero_config(tmp_path: Path) -> None:
    table, template = _fixtures(tmp_path)
    out = tmp_path / 'out'
    result = CliRunner().invoke(
        app,
        [
            'generate',
            '--table', str(table),
            '--template', str(template),
            '--out', str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert len(list(out.glob('*.docx'))) == 2


def test_cli_generate_reports_error_on_header_collision(tmp_path: Path) -> None:
    headers = ['Дата начала', 'Дата_начала']
    table = _make_xlsx(tmp_path / 'emp.xlsx', [['a', 'b']], headers)
    template = _make_template(tmp_path / 'tpl.docx')
    result = CliRunner().invoke(
        app,
        [
            'generate',
            '--table', str(table),
            '--template', str(template),
            '--out', str(tmp_path / 'out'),
        ],
    )
    assert result.exit_code == 1


def test_gender_overrides_parses_and_normalizes() -> None:
    cfg = Config.model_validate({'genders': {'Ким Саша': 'ж', 'Иванов Пётр': 'м'}})
    overrides = _gender_overrides(cfg)
    assert overrides == {'ким саша': Gender.FEMALE, 'иванов пётр': Gender.MALE}


def test_gender_overrides_warns_on_unparseable_value(
    caplog: pytest.LogCaptureFixture,
) -> None:
    cfg = Config.model_validate({'genders': {'Ким Саша': 'мужык'}})
    with caplog.at_level(logging.WARNING):
        overrides = _gender_overrides(cfg)
    assert overrides == {}
    assert any('мужык' in r.message for r in caplog.records)


def test_cli_check_clean_table_exits_zero(tmp_path: Path) -> None:
    table = _make_xlsx(tmp_path / 'emp.xlsx', _ROWS, _HEADERS)
    template = _make_template(tmp_path / 'tpl.docx')
    result = CliRunner().invoke(
        app,
        [
            'check',
            '--table', str(table),
            '--template', str(template),
        ],
    )
    assert result.exit_code == 0, result.output
    assert 'Проблем не найдено' in result.output


def test_cli_check_undefined_var_exits_one(tmp_path: Path) -> None:
    table = _make_xlsx(tmp_path / 'emp.xlsx', _ROWS, _HEADERS)
    template = tmp_path / 'bad.docx'
    doc = Document()
    doc.add_paragraph('{{ Опечатка }}')
    doc.save(template)
    result = CliRunner().invoke(
        app,
        [
            'check',
            '--table', str(table),
            '--template', str(template),
        ],
    )
    assert result.exit_code == 1
    assert 'Опечатка' in result.output


class _RecordingStream:
    def __init__(self) -> None:
        self.encodings: list[str] = []

    def reconfigure(self, *, encoding: str) -> None:
        self.encodings.append(encoding)


def test_configure_stdio_switches_to_utf8_when_flagged(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # T023: ядро из GUI (PYTHONUTF8=1) переключает потоки на UTF-8.
    monkeypatch.setenv('PYTHONUTF8', '1')
    out, err = _RecordingStream(), _RecordingStream()
    monkeypatch.setattr('dyak.cli.sys.stdout', out)
    monkeypatch.setattr('dyak.cli.sys.stderr', err)
    configure_stdio()
    assert out.encodings == ['utf-8']
    assert err.encodings == ['utf-8']


def test_configure_stdio_noop_without_flag(monkeypatch: pytest.MonkeyPatch) -> None:
    # Прямой CLI (без PYTHONUTF8) — потоки не трогаем (нативная консоль).
    monkeypatch.delenv('PYTHONUTF8', raising=False)
    out = _RecordingStream()
    monkeypatch.setattr('dyak.cli.sys.stdout', out)
    configure_stdio()
    assert out.encodings == []
